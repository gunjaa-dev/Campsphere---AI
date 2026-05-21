"""
CAMSPHER-AI Resume Analyzer — Multi-stage text extraction v2.0

Supports: PDF (text + scanned), DOCX, DOC, images (JPG/PNG/WEBP).
Pipeline: native extractors → PyMuPDF → OCR (pytesseract + EasyOCR).
Never hard-fails; always returns best-effort text with metadata.
"""

from __future__ import annotations

import io
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from typing import Any, Dict, List, Optional, Tuple

from utils.ocr_engine import ocr_bytes, ocr_image, ocr_images

logger = logging.getLogger(__name__)

MIN_TEXT_THRESHOLD = 250
MIN_CHARS_ACCEPT = 120

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
DOC_EXTENSIONS = {".doc"}

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


class TextExtractor:
    """Multi-stage resume text extraction with OCR fallbacks."""

    def __init__(self):
        self.supported_formats = sorted(
            PDF_EXTENSIONS | DOCX_EXTENSIONS | DOC_EXTENSIONS | IMAGE_EXTENSIONS
        )

    def detect_file_type(
        self,
        filename: str,
        file_content: bytes,
        content_type: Optional[str] = None,
    ) -> str:
        ext = os.path.splitext((filename or "").lower())[1]
        if ext in IMAGE_EXTENSIONS:
            return "image"
        if ext in PDF_EXTENSIONS:
            return "pdf"
        if ext in DOCX_EXTENSIONS:
            return "docx"
        if ext in DOC_EXTENSIONS:
            return "doc"

        mime = content_type or mimetypes.guess_type(filename or "")[0] or ""
        if mime.startswith("image/"):
            return "image"
        if mime == "application/pdf" or file_content[:4] == b"%PDF":
            return "pdf"
        if mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
        ) or file_content[:2] == b"PK":
            return "docx"
        if mime == "application/msword" or file_content[:4] == b"\xd0\xcf\x11\xe0":
            return "doc"

        return "unknown"

    def extract(self, file_content: bytes, filename: str, content_type: Optional[str] = None) -> str:
        """Backward-compatible: returns cleaned text only."""
        result = self.extract_with_metadata(file_content, filename, content_type)
        return result["cleaned_text"]

    def extract_with_metadata(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None,
    ) -> Dict[str, Any]:
        file_type = self.detect_file_type(filename, file_content, content_type)
        warnings: List[str] = []

        if file_type == "pdf":
            raw, method, conf = self._extract_pdf_pipeline(file_content, warnings)
        elif file_type == "docx":
            raw, method, conf = self._extract_docx_pipeline(file_content, warnings)
        elif file_type == "doc":
            raw, method, conf = self._extract_doc_pipeline(file_content, filename, warnings)
        elif file_type == "image":
            raw, method, conf = self._extract_image_pipeline(file_content, warnings)
        else:
            raw, method, conf = self._extract_unknown_pipeline(file_content, filename, warnings)

        cleaned = self._clean_text(raw)

        if len(cleaned) < MIN_TEXT_THRESHOLD:
            warnings.append("low_text_yield_retry")
            retry_raw, retry_method, retry_conf = self._retry_aggressive(
                file_content, filename, file_type
            )
            if len(retry_raw) > len(raw):
                raw = retry_raw
                method = f"{method}+{retry_method}" if method else retry_method
                conf = max(conf, retry_conf)
                cleaned = self._clean_text(raw)

        if len(cleaned) < MIN_CHARS_ACCEPT:
            warnings.append(
                "Very little text extracted; analysis scores may be limited. "
                "Ensure the file is a clear resume scan or photo."
            )
            if not cleaned.strip():
                cleaned = self._minimal_fallback_text(filename)

        return {
            "raw_text": raw,
            "cleaned_text": cleaned,
            "extraction_method_used": method or "fallback",
            "confidence_score": round(min(100.0, max(0.0, conf)), 1),
            "file_type_detected": file_type,
            "warnings": warnings,
            "character_count": len(cleaned),
        }

    def _minimal_fallback_text(self, filename: str) -> str:
        base = os.path.splitext(os.path.basename(filename or "resume"))[0]
        base = re.sub(r"[_\-]+", " ", base)
        return (
            f"{base}\n"
            "Professional Summary Experience Education Skills Projects Certifications\n"
            "Work History Technical Skills Soft Skills"
        )

    def _retry_aggressive(
        self, file_content: bytes, filename: str, file_type: str
    ) -> Tuple[str, str, float]:
        if file_type == "pdf":
            return self._pdf_ocr_pipeline(file_content, aggressive=True)
        if file_type == "image":
            text, conf, method = ocr_bytes(file_content, aggressive=True)
            return text, method, conf
        if file_type == "docx":
            return self._extract_docx_pipeline(file_content, [])
        if file_type == "doc":
            return self._extract_doc_pipeline(file_content, filename, [])
        return "", "retry_none", 0.0

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _extract_pdf_pipeline(
        self, file_content: bytes, warnings: List[str]
    ) -> Tuple[str, str, float]:
        methods: List[Tuple[str, str, float]] = []

        t, m, c = self._pdf_pdfplumber(file_content)
        if len(t) >= MIN_TEXT_THRESHOLD:
            return t, m, c
        if t:
            methods.append((t, m, c))

        t, m, c = self._pdf_pymupdf_text(file_content)
        if len(t) >= MIN_TEXT_THRESHOLD:
            return t, m, c
        if t:
            methods.append((t, m, c))

        if PDFMINER_AVAILABLE:
            try:
                laparams = LAParams(line_margin=0.5, word_margin=0.1, char_margin=2.0)
                t = pdfminer_extract(io.BytesIO(file_content), laparams=laparams) or ""
                if len(t.strip()) >= MIN_TEXT_THRESHOLD:
                    return t.strip(), "pdfminer", 75.0
                if t.strip():
                    methods.append((t.strip(), "pdfminer", 70.0))
            except Exception as e:
                warnings.append(f"pdfminer: {e}")

        warnings.append("pdf_native_low_yield_ocr")
        ocr_text, ocr_method, ocr_conf = self._pdf_ocr_pipeline(file_content, aggressive=False)
        if ocr_text:
            if methods:
                combined = self._merge_text_parts([m[0] for m in methods] + [ocr_text])
                best_conf = max([m[2] for m in methods] + [ocr_conf])
                return combined, f"{methods[-1][1]}+{ocr_method}", best_conf
            return ocr_text, ocr_method, ocr_conf

        if methods:
            combined = self._merge_text_parts([m[0] for m in methods])
            return combined, "+".join(m[1] for m in methods), max(m[2] for m in methods)

        return ocr_text, ocr_method, ocr_conf

    def _pdf_pdfplumber(self, file_content: bytes) -> Tuple[str, str, float]:
        if not PDFPLUMBER_AVAILABLE:
            return "", "pdfplumber_unavailable", 0.0
        try:
            pages = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    pt = page.extract_text()
                    if pt and pt.strip():
                        pages.append(pt)
            text = "\n\n".join(pages)
            conf = min(95.0, 50.0 + len(text) / 50.0) if text else 0.0
            return text, "pdfplumber", conf
        except Exception as e:
            logger.debug("pdfplumber: %s", e)
            return "", "pdfplumber_failed", 0.0

    def _pdf_pymupdf_text(self, file_content: bytes) -> Tuple[str, str, float]:
        if not PYMUPDF_AVAILABLE:
            return "", "pymupdf_unavailable", 0.0
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            pages = []
            for page in doc:
                pt = page.get_text("text") or page.get_text()
                if pt and pt.strip():
                    pages.append(pt)
            doc.close()
            text = "\n\n".join(pages)
            conf = min(92.0, 48.0 + len(text) / 55.0) if text else 0.0
            return text, "pymupdf", conf
        except Exception as e:
            logger.debug("pymupdf text: %s", e)
            return "", "pymupdf_failed", 0.0

    def _pdf_pages_to_images_pymupdf(self, file_content: bytes, dpi: int = 200) -> List["Image.Image"]:
        if not PYMUPDF_AVAILABLE or not PIL_AVAILABLE:
            return []
        images = []
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            for page in doc:
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                images.append(img)
            doc.close()
        except Exception as e:
            logger.debug("pymupdf render: %s", e)
        return images

    def _pdf_pages_to_images_pdf2image(self, file_content: bytes, dpi: int = 200) -> List["Image.Image"]:
        if not PDF2IMAGE_AVAILABLE or not PIL_AVAILABLE:
            return []
        try:
            return convert_from_bytes(file_content, dpi=dpi, fmt="png")
        except Exception as e:
            logger.debug("pdf2image: %s", e)
            return []

    def _pdf_ocr_pipeline(
        self, file_content: bytes, aggressive: bool = False
    ) -> Tuple[str, str, float]:
        dpi = 300 if aggressive else 200
        images = self._pdf_pages_to_images_pymupdf(file_content, dpi=dpi)
        method_prefix = "pymupdf_render"
        if not images:
            images = self._pdf_pages_to_images_pdf2image(file_content, dpi=dpi)
            method_prefix = "pdf2image"

        if not images:
            return "", f"{method_prefix}_ocr_failed", 0.0

        text, conf, ocr_method = ocr_images(images, aggressive=aggressive)
        return text, f"{method_prefix}+{ocr_method}", conf

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _extract_docx_pipeline(
        self,
        file_content: bytes,
        warnings: List[str],
        called_from_fallback: bool=False
        ) -> Tuple[str,str,float]:

        if not PYTHON_DOCX_AVAILABLE:
            warnings.append("python-docx unavailable")

            if called_from_fallback:
                return "", "docx_failed", 0.0

            return self._extract_doc_as_zip_fallback(
                file_content,
                warnings
            )

        try:
            doc = Document(io.BytesIO(file_content))

            parts=[]

            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt=cell.text.strip()
                        if txt:
                            parts.append(txt)

            text="\n".join(parts)

            conf=min(
                95.0,
                60+(len(text)/40)
            ) if text else 0

            return text,"python-docx",conf

        except Exception as e:

            logger.debug(
                f"DOCX parse failed: {e}"
            )

            if called_from_fallback:
                return "", "docx_failed", 0.0

            return self._extract_doc_as_zip_fallback(
                file_content,
                warnings
            )


    def _extract_doc_as_zip_fallback(
        self,
        file_content: bytes,
        warnings: List[str]=None
    )->Tuple[str,str,float]:

        if warnings is None:
            warnings=[]

        import zipfile

        try:

            z=zipfile.ZipFile(
                io.BytesIO(file_content)
            )

            xml=z.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="ignore"
            )

            text=re.sub(
                "<[^>]+>",
                " ",
                xml
            )

            text=self._clean_text(text)

            if len(text)>50:
                return (
                    text,
                    "docx_zip_fallback",
                    70.0
                )

        except Exception as e:

            logger.debug(
                f"ZIP fallback failed:{e}"
            )

        return "", "docx_failed", 0.0

    # ── DOC ───────────────────────────────────────────────────────────────────

    def _extract_doc_pipeline(
        self, file_content: bytes, filename: str, warnings: List[str]
    ) -> Tuple[str, str, float]:
        if file_content[:2] == b"PK":
            return self._extract_docx_pipeline(file_content, warnings)

        if TEXTRACT_AVAILABLE:
            try:
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                    tmp.write(file_content)
                    tmp.flush()
                    raw = textract.process(tmp.name).decode("utf-8", errors="ignore")
                    os.unlink(tmp.name)
                if raw.strip():
                    return raw.strip(), "textract", 80.0
            except Exception as e:
                warnings.append(f"textract: {e}")

        antiword_text = self._antiword_extract(file_content, filename)
        if antiword_text:
            return antiword_text, "antiword", 75.0

        warnings.append("doc_legacy_weak_try_ocr")
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=file_content, filetype="doc")
                pages = [page.get_text() for page in doc if page.get_text()]
                doc.close()
                text = "\n".join(pages)
                if text.strip():
                    return text.strip(), "pymupdf_doc", 65.0
            except Exception:
                pass

        return "", "doc_unsupported_format", 25.0

    def _antiword_extract(self, file_content: bytes, filename: str) -> str:
        antiword = shutil.which("antiword")
        if not antiword:
            return ""
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(file_content)
                tmp.flush()
                out = subprocess.run(
                    [antiword, tmp.name],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                os.unlink(tmp.name)
            if out.returncode == 0 and out.stdout:
                return out.stdout.strip()
        except Exception as e:
            logger.debug("antiword: %s", e)
        return ""

    # ── Images ────────────────────────────────────────────────────────────────

    def _extract_image_pipeline(
        self, file_content: bytes, warnings: List[str]
    ) -> Tuple[str, str, float]:
        if not PIL_AVAILABLE:
            warnings.append("Pillow not installed for image OCR")
            text, conf, method = ocr_bytes(file_content, aggressive=False)
            return text, method, conf

        try:
            img = Image.open(io.BytesIO(file_content))
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            text, conf, method = ocr_image(img, aggressive=False)
            return text, f"image_{method}", conf
        except Exception as e:
            warnings.append(f"image_open: {e}")
            return ocr_bytes(file_content, aggressive=False)

    def _extract_unknown_pipeline(
        self, file_content: bytes, filename: str, warnings: List[str]
    ) -> Tuple[str, str, float]:
        warnings.append("unknown_type_trying_all")
        for fn in (
            lambda: self._extract_pdf_pipeline(file_content, warnings),
            lambda: self._extract_docx_pipeline(file_content, warnings),
            lambda: self._extract_image_pipeline(file_content, warnings),
        ):
            try:
                t, m, c = fn()
                if len(t) >= MIN_CHARS_ACCEPT:
                    return t, f"auto_{m}", c
            except Exception:
                continue
        return "", "unknown", 0.0

    @staticmethod
    def _merge_text_parts(parts: List[str]) -> str:
        seen = set()
        lines = []
        for part in parts:
            for line in part.splitlines():
                key = line.strip().lower()
                if key and key not in seen:
                    seen.add(key)
                    lines.append(line.strip())
        return "\n".join(lines)

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""

        text = text.replace("\x00", "")
        text = re.sub(r"[\u200b-\u200f\u202a-\u202e\ufeff]", "", text)
        text = re.sub(r"[^\S\n]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        text = "".join(c for c in text if c.isprintable() or c in "\n\t")

        lines = []
        for line in text.split("\n"):
            line = line.strip()
            line = re.sub(r"^[^\w\s]{1,3}$", "", line)
            if len(line) > 1 or line == "":
                lines.append(line)
        text = "\n".join(lines)

        text = re.sub(
            r"(.)\1{4,}",
            lambda m: m.group(1) * 2,
            text,
        )
        return text.strip()

    def extract_from_text(self, text: str) -> str:
        return self._clean_text(text)


def extract_text(file_content: bytes, filename: str, content_type: Optional[str] = None) -> str:
    return TextExtractor().extract(file_content, filename, content_type)


def extract_text_with_metadata(
    file_content: bytes, filename: str, content_type: Optional[str] = None
) -> Dict[str, Any]:
    return TextExtractor().extract_with_metadata(file_content, filename, content_type)
